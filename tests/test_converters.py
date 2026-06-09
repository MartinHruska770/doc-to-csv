import io

import pytest
from docx import Document

from converters import (
    UnsupportedFormat,
    extract_lines,
    sentences_to_csv,
    split_sentences,
)


def _docx_bytes(build) -> io.BytesIO:
    document = Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def test_split_sentences_basic():
    text = "První věta. Druhá věta! Třetí věta?"
    assert list(split_sentences([text])) == [
        "První věta.",
        "Druhá věta!",
        "Třetí věta?",
    ]


def test_split_sentences_joins_paragraphs():
    paragraphs = ["První věta. Druhá věta.", "Třetí věta."]
    assert list(split_sentences(paragraphs)) == [
        "První věta.",
        "Druhá věta.",
        "Třetí věta.",
    ]


def test_split_sentences_ignores_blank_lines():
    assert list(split_sentences(["", "   ", "Jediná věta."])) == ["Jediná věta."]


def test_split_sentences_keeps_text_without_terminator():
    assert list(split_sentences(["Bez tečky na konci"])) == ["Bez tečky na konci"]


def test_split_sentences_does_not_break_on_title_abbreviation():
    text = "Přišel Mgr. Novák. Pozdravil nás."
    assert list(split_sentences([text])) == [
        "Přišel Mgr. Novák.",
        "Pozdravil nás.",
    ]


def test_split_sentences_handles_common_abbreviation():
    text = "Vezmi ovoce, např. jablka. Pak je umyj."
    assert list(split_sentences([text])) == [
        "Vezmi ovoce, např. jablka.",
        "Pak je umyj.",
    ]


def test_split_sentences_does_not_break_on_date():
    text = "Sešli jsme se 1. ledna 2024. Bylo zima."
    assert list(split_sentences([text])) == [
        "Sešli jsme se 1. ledna 2024.",
        "Bylo zima.",
    ]


def test_split_sentences_does_not_break_on_ordinal_with_capital():
    text = "Probíhala 2. Světová válka. Konec přišel v roce 1945."
    assert list(split_sentences([text])) == [
        "Probíhala 2. Světová válka.",
        "Konec přišel v roce 1945.",
    ]


def test_split_sentences_multiple_abbreviations_in_row():
    text = "Pozval Ing. Nováka a doc. Svobodu. Oba dorazili."
    assert list(split_sentences([text])) == [
        "Pozval Ing. Nováka a doc. Svobodu.",
        "Oba dorazili.",
    ]


def test_split_sentences_phd_title():
    text = "Mluvil Petr Novák, Ph.D. Druhý přednášející byl jiný."
    assert list(split_sentences([text])) == [
        "Mluvil Petr Novák, Ph.D. Druhý přednášející byl jiný.",
    ]


def test_sentences_to_csv_one_per_row():
    csv_text = sentences_to_csv(["První.", "Druhá."])
    assert csv_text.splitlines() == ["﻿První.", "Druhá."]


def test_sentences_to_csv_quotes_when_needed():
    csv_text = sentences_to_csv(["Věta, s čárkou."])
    assert csv_text.splitlines() == ["﻿\"Věta, s čárkou.\""]


def test_sentences_to_csv_empty():
    assert sentences_to_csv([]) == ""


def test_unknown_extension_raises():
    with pytest.raises(UnsupportedFormat):
        list(extract_lines(io.BytesIO(b""), ".odt"))


def test_extract_txt_plain():
    data = "První věta. Druhá věta.".encode("utf-8")
    sentences = list(split_sentences(extract_lines(io.BytesIO(data), ".txt")))
    assert sentences == ["První věta.", "Druhá věta."]


def test_extract_txt_with_bom():
    data = "﻿První věta.".encode("utf-8")
    sentences = list(split_sentences(extract_lines(io.BytesIO(data), ".txt")))
    assert "První věta." in sentences


def test_extract_docx_reads_paragraphs_and_tables():
    def build(doc):
        doc.add_paragraph("Nadpis dokumentu.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Buňka A."
        table.rows[0].cells[1].text = "Buňka B."

    sentences = list(split_sentences(extract_lines(_docx_bytes(build), ".docx")))
    assert "Nadpis dokumentu." in sentences
    assert "Buňka A." in sentences
    assert "Buňka B." in sentences
