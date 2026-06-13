import csv
import io
import re
from collections.abc import Iterable, Iterator
from typing import BinaryIO

import pdfplumber
from docx import Document

BOM = "﻿"

_BOUNDARY = re.compile(r"(?P<punct>[.!?]+)\s+(?=[A-ZÁ-Ž0-9„\"'(])")
_LAST_TOKEN = re.compile(r"([A-Za-zÁ-Žá-ž]+|\d+)$")
_DOTTED_ACRONYM = re.compile(r"[A-Za-zÁ-Žá-ž]\.[A-Za-zÁ-Žá-ž]{1,3}$")

_ABBREVIATIONS = {
    "mgr", "bc", "ing", "ph", "phd", "csc", "drsc",
    "mudr", "judr", "phdr", "rndr", "mvdr", "paeddr",
    "prof", "doc", "dr",
    "např", "atd", "atp", "apod", "aj", "tj", "tzv", "tzn",
    "resp", "popř", "příp", "kupř", "zvl", "zejm",
    "vč", "mj", "cca", "max", "min", "event", "ev",
    "str", "s", "č", "kap", "odst", "písm", "obr", "tab",
    "roč", "sv", "vyd", "ed", "pozn",
    "n", "l", "p", "pí", "sl", "stol",
}


class UnsupportedFormat(ValueError):
    ...


def _from_pdf(stream):
    """Vydává řádky (de-hyphenované) a prázdné řetězce na hranicích odstavců."""
    with pdfplumber.open(stream) as pdf:
        for page in pdf.pages:
            pending = None
            for raw in (page.extract_text(layout=True) or "").splitlines():
                stripped = raw.strip()
                if not stripped:
                    # prázdný řádek = hranice odstavce
                    if pending is not None:
                        yield pending
                        pending = None
                    yield ""
                elif pending is not None and pending.endswith("-") and pending[-2:-1].islower():
                    pending = pending[:-1] + stripped  # de-hyphenace
                else:
                    if pending is not None:
                        yield pending
                    pending = stripped
            if pending is not None:
                yield pending
                pending = None
            yield ""  # konec stránky = hranice odstavce


def _from_docx(stream):
    document = Document(stream)
    for paragraph in document.paragraphs:
        yield paragraph.text
    for table in document.tables:
        for row in table.rows:
            yield " ".join(cell.text for cell in row.cells)


def _from_txt(stream):
    yield from stream.read().decode("utf-8-sig").splitlines()


_EXTRACTORS = {
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".txt": _from_txt,
}

SUPPORTED_EXTENSIONS = set(_EXTRACTORS)


def extract_lines(stream: BinaryIO, extension: str) -> Iterator[str]:
    try:
        extractor = _EXTRACTORS[extension.lower()]
    except KeyError as exc:
        raise UnsupportedFormat(extension) from exc
    return extractor(stream)


def _paragraphs_from_lines(lines: Iterable[str]) -> Iterator[str]:
    """Seskupí řádky do odstavců – předěl tvoří prázdný řádek."""
    buffer: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            buffer.append(stripped)
        elif buffer:
            yield " ".join(buffer)
            buffer = []
    if buffer:
        yield " ".join(buffer)


def _paragraphs_from_docx(stream):
    document = Document(stream)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in document.tables:
        for row in table.rows:
            text = " ".join(cell.text for cell in row.cells).strip()
            if text:
                yield text


_PARAGRAPH_EXTRACTORS = {
    ".pdf": lambda stream: _paragraphs_from_lines(_from_pdf(stream)),
    ".docx": _paragraphs_from_docx,
    ".txt": lambda stream: _paragraphs_from_lines(_from_txt(stream)),
}


def extract_paragraphs(stream: BinaryIO, extension: str) -> Iterator[str]:
    """Vrací jednotlivé odstavce dokumentu (jednotka pro anotaci metafor)."""
    try:
        extractor = _PARAGRAPH_EXTRACTORS[extension.lower()]
    except KeyError as exc:
        raise UnsupportedFormat(extension) from exc
    return extractor(stream)


def _is_real_boundary(text_before):
    if _DOTTED_ACRONYM.search(text_before):
        return False
    token = _LAST_TOKEN.search(text_before)
    if token is None:
        return True
    value = token.group(0)
    if value.isdigit():
        return len(value) > 3
    return value.casefold() not in _ABBREVIATIONS


def split_sentences(lines: Iterable[str]) -> Iterator[str]:
    text = " ".join(line.strip() for line in lines if line and line.strip())
    if not text:
        return

    cursor = 0
    for match in _BOUNDARY.finditer(text):
        if not _is_real_boundary(text[cursor:match.start()]):
            continue
        sentence = text[cursor:match.end("punct")].strip()
        if sentence:
            yield sentence
        cursor = match.end()

    tail = text[cursor:].strip()
    if tail:
        yield tail


def sentences_to_csv(sentences: Iterable[str]) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    for sentence in sentences:
        writer.writerow([sentence])
    data = buffer.getvalue()
    return BOM + data if data else ""


def units_to_csv(units: Iterable[str], prefix: str = "") -> str:
    """CSV pro anotaci: hlavička ``id,text`` a jeden očíslovaný odstavec na řádek.

    ``prefix`` umožní jednoznačné ID i po sloučení více dokumentů
    (``zprava-1``, ``zprava-2`` …); bez něj jsou ID prostá čísla.
    """
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    count = 0
    for unit in units:
        text = unit.strip()
        if not text:
            continue
        count += 1
        identifier = f"{prefix}-{count}" if prefix else str(count)
        writer.writerow([identifier, text])
    if not count:
        return ""
    header = io.StringIO()
    csv.writer(header).writerow(["id", "text"])
    return BOM + header.getvalue() + buffer.getvalue()
